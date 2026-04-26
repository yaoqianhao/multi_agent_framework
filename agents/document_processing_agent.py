"""
文档处理智能体 - Document Processing Agent

这是一个可以集成到多智能体框架中的智能体，专门用于读取、解析、创建和修改
PDF、Word (DOCX) 以及其他常见文本文档的内容。

依赖安装:
    pip install pypdf pdfplumber python-docx

使用示例:
    from agents.document_processing_agent import create_document_processing_agent
    from multi_agent_framework import AgentFramework
    
    framework = AgentFramework()
    doc_agent = create_document_processing_agent()
    framework.register_agent(doc_agent)
    
    result = framework.execute(
        task="读取 sample.pdf 中的文本并提取出核心观点保存到 result.docx",
        workflow_type="pipeline",
        workflow_config={"agent_sequence": ["document_processor"]}
    )
"""

from __future__ import annotations

import os
from typing import Any, Dict, List, Optional
from langchain_core.tools import tool
from multi_agent_framework import ReactAgent, BaseAgent, get_llm


# =============================================================================
# PDF 处理工具
# =============================================================================

@tool
def extract_pdf_text(file_path: str, max_pages: int = 0) -> str:
    """
    读取并提取 PDF 文件中的纯文本内容。
    
    Args:
        file_path: PDF文件的完整绝对路径
        max_pages: 最大读取页数（0表示读取所有页）
        
    Returns:
        提取出的PDF文本内容
    """
    try:
        from pypdf import PdfReader
        
        if not os.path.exists(file_path):
            return f"错误: 找不到文件 {file_path}"
            
        reader = PdfReader(file_path)
        num_pages = len(reader.pages)
        
        pages_to_read = num_pages if max_pages <= 0 else min(num_pages, max_pages)
        
        text = f"--- PDF 文件 '{os.path.basename(file_path)}' 内容 (共 {num_pages} 页, 读取前 {pages_to_read} 页) ---\n\n"
        
        for i in range(pages_to_read):
            page_text = reader.pages[i].extract_text()
            if page_text:
                text += f"[第 {i+1} 页]\n{page_text}\n\n"
                
        return text
    except Exception as e:
        return f"读取PDF时发生错误: {str(e)}"

@tool
def merge_pdf_files(input_paths: List[str], output_path: str) -> str:
    """
    将多个 PDF 文件按顺序合并为一个新的 PDF 文件。
    
    Args:
        input_paths: 要合并的 PDF 文件绝对路径列表
        output_path: 合并后生成的 PDF 文件绝对路径
        
    Returns:
        操作结果消息
    """
    try:
        from pypdf import PdfWriter, PdfReader
        
        writer = PdfWriter()
        for path in input_paths:
            if not os.path.exists(path):
                return f"错误: 找不到输入文件 {path}"
            
            reader = PdfReader(path)
            for page in reader.pages:
                writer.add_page(page)
                
        with open(output_path, "wb") as f:
            writer.write(f)
            
        return f"成功: 已将 {len(input_paths)} 个文件合并并保存至 {output_path}"
    except Exception as e:
        return f"合并PDF时发生错误: {str(e)}"


# =============================================================================
# Word (DOCX) 处理工具
# =============================================================================

@tool
def extract_docx_text(file_path: str) -> str:
    """
    读取并提取 Word (DOCX) 文件中的文本内容。
    
    Args:
        file_path: DOCX文件的完整绝对路径
        
    Returns:
        提取出的文本内容
    """
    try:
        from docx import Document
        
        if not os.path.exists(file_path):
            return f"错误: 找不到文件 {file_path}"
            
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
            
        return "\n".join(text)
    except Exception as e:
        return f"读取DOCX时发生错误: {str(e)}"

@tool
def create_docx_document(file_path: str, content: str, title: Optional[str] = None) -> str:
    """
    创建一个新的 Word (DOCX) 文档，并写入指定内容。
    
    Args:
        file_path: 要保存的DOCX文件的完整绝对路径
        content: 要写入正文的文本内容（使用换行符分隔段落）
        title: 文档标题（可选）
        
    Returns:
        操作结果消息
    """
    try:
        from docx import Document
        
        doc = Document()
        
        if title:
            doc.add_heading(title, 0)
            
        # 根据换行符分割段落
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():  # 忽略空行
                doc.add_paragraph(para)
                
        # 确保目录存在
        os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
        doc.save(file_path)
        
        return f"成功: 已创建文档并保存至 {file_path}"
    except Exception as e:
        return f"创建DOCX时发生错误: {str(e)}"

@tool
def replace_text_in_docx(file_path: str, old_text: str, new_text: str, output_path: Optional[str] = None) -> str:
    """
    在 Word (DOCX) 文件中查找并替换指定的文本，可以保存为新文件或覆盖原文件。
    注意：此操作替换段落中的文本，可能无法覆盖页眉/页脚/表格中的特殊格式。
    
    Args:
        file_path: 源DOCX文件的完整绝对路径
        old_text: 要被替换的旧文本
        new_text: 替换成的新文本
        output_path: 替换后保存的文件路径。如果为空则覆盖原文件
        
    Returns:
        操作结果消息及替换次数
    """
    try:
        from docx import Document
        
        if not os.path.exists(file_path):
            return f"错误: 找不到源文件 {file_path}"
            
        doc = Document(file_path)
        replace_count = 0
        
        # 遍历段落中的 runs，保留原有的字体格式
        for para in doc.paragraphs:
            if old_text in para.text:
                # 简单替换：直接对 text 赋值会丢失 run 级别格式。这里做全局段落文本替换
                # 对于更精细的带格式保留替换，需要极其复杂的 run 合并逻辑，这里暂用全段替换。
                inline = para.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(old_text, new_text)
                        replace_count += 1
                        
        # 补充：替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            inline = para.runs
                            for i in range(len(inline)):
                                if old_text in inline[i].text:
                                    inline[i].text = inline[i].text.replace(old_text, new_text)
                                    replace_count += 1
        
        save_path = output_path if output_path else file_path
        doc.save(save_path)
        
        return f"成功: 共替换了 {replace_count} 处文本。文件已保存至 {save_path}"
    except Exception as e:
        return f"修改DOCX时发生错误: {str(e)}"


# =============================================================================
# 纯文本处理工具
# =============================================================================

@tool
def read_text_file(file_path: str) -> str:
    """
    读取普通的纯文本文件 (.txt, .md, .csv, 等等)。
    
    Args:
        file_path: 文本文件的完整绝对路径
        
    Returns:
        文件内容
    """
    try:
        if not os.path.exists(file_path):
            return f"错误: 找不到文件 {file_path}"
            
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"读取文本文件时发生错误: {str(e)}"

@tool
def write_text_file(file_path: str, content: str) -> str:
    """
    写入内容到纯文本文件（如果文件存在则覆盖）。
    
    Args:
        file_path: 文本文件的完整绝对路径
        content: 要写入的内容
        
    Returns:
        操作结果消息
    """
    try:
        os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"成功: 文件内容已写入 {file_path}"
    except Exception as e:
        return f"写入文本文件时发生错误: {str(e)}"


# =============================================================================
# 智能体工厂方法
# =============================================================================

def create_document_processing_agent() -> BaseAgent:
    """
    创建并返回文档处理智能体。
    """
    
    system_prompt = (
        "你是一个高级文档处理助手 (Document Processing Agent)。\n"
        "你的主要任务是读取、修改、提取和生成各种类型的文档内容（包括 PDF、DOCX、TXT 等）。\n"
        "你可以使用以下工具：\n"
        "- extract_pdf_text: 读取 PDF 文件。\n"
        "- merge_pdf_files: 合并多个 PDF。\n"
        "- extract_docx_text: 读取 Word 文件。\n"
        "- create_docx_document: 创建新的 Word 文件。\n"
        "- replace_text_in_docx: 替换 Word 文件中的特定文本。\n"
        "- read_text_file / write_text_file: 读写普通文本文档。\n\n"
        "工作流程准则：\n"
        "1. 在修改文档前，通常先阅读文档了解当前内容。\n"
        "2. 执行文件操作后，务必将生成或修改的新文件路径告知用户。\n"
        "3. 在处理报错时，仔细分析错误原因，比如文件不存在或格式不正确，并主动修复。\n"
        "4. 操作路径必须使用绝对路径以防止找不到文件。"
    )
    
    tools = [
        extract_pdf_text,
        merge_pdf_files,
        extract_docx_text,
        create_docx_document,
        replace_text_in_docx,
        read_text_file,
        write_text_file
    ]
    
    # 我们使用 ReactAgent 来赋予它调用工具和自主思考的能力
    return ReactAgent(
        name="document_processor",
        description="专门负责处理和修改 PDF、Word (DOCX) 等文档内容的智能体",
        system_prompt=system_prompt,
        temperature=0.1,  # 降低随机性，确保工具调用的精准性
        tools=tools
    )

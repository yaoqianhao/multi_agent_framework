import os
from agents.document_processing_agent import create_document_processing_agent
from multi_agent_framework import AgentFramework

def main():
    print("初始化 Agent Framework 和 Document Processing Agent...")
    # 1. 注册智能体
    framework = AgentFramework()
    doc_agent = create_document_processing_agent()
    framework.register_agent(doc_agent)

    # 请确保您的 E 盘根目录下有一个叫做 '合同.docx' 的文件
    # 或者您可以修改下面的 task 以匹配您实际的文件路径
    target_file = "E:/合同.docx"
    
    if not os.path.exists(target_file):
        print(f"\n[警告] 找不到文件 {target_file}")
        print("请创建一个测试用的Word文件，或修改代码中的文件路径。")
        # 自动生成一个测试文件以防报错
        print("正在自动为您生成一个测试用的 E:/合同.docx...")
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("这是一份测试合同。")
            doc.add_paragraph("甲方：张三")
            doc.add_paragraph("乙方：李四")
            doc.save(target_file)
            print(f"成功生成测试文件：{target_file}")
        except Exception as e:
            print(f"自动生成测试文件失败: {e}")
            return

    print("\n发起任务：替换 word 里的内容")
    # 2. 发起任务（比如替换 word 里的内容，或者从 PDF 中提取总结）
    result = framework.execute(
        task=f"帮我把 '{target_file}' 里的 '甲方' 全部替换成 'CodeBuddy有限公司'",
        workflow_type="pipeline",
        workflow_config={"agent_sequence": ["document_processor"]}
    )
    
    print("\n" + "="*50)
    print("任务执行结果：")
    print(result.get("result", "无返回结果"))
    print("="*50)

if __name__ == "__main__":
    main()
